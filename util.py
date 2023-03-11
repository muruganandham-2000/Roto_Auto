import pandas as pd
import pyodbc
import mysql.connector as db
from docx.shared import Inches
import docx
import os
import shutil

def Db_Connection():
    conn=db.connect(host='localhost',user='root',password='Muruga@2000')
    if conn.is_connected():
        print("DB Connected...")
    else:
        print("connection not established...")
    cursor=conn.cursor()
    cursor.execute("SELECT * FROM pydatabase.py_automate where Run_Indicator='Y';")
    record = cursor.fetchall()
    cursor.close()
    conn.close()
    return record

def input_data(data):
    Scenario,F_Name,L_Name,Email,Comment = ([] for i in range(5))
    for i in range (len(data)):
        Scenario.append(data[i][0])
        F_Name.append(data[i][2])
        L_Name.append(data[i][3])
        Email.append(data[i][4])
        Comment.append(data[i][5])
    mydict={"SCENARIO":Scenario,"F_NAME":F_Name,"L_NAME":L_Name,"EMAIL":Email,"COMMENT":Comment}
    return mydict

mylist=[]
def Screenshot(driver,SS_msg):
        try:
            mylist.append(SS_msg)
            directory = "temp"
            parent_dir = "D:\End to End\Evidence"
            path = os.path.join(parent_dir,directory)
            if os.path.isdir(path)==False:
                os.mkdir(path)
            driver.save_screenshot('D:\End to End\Evidence\\'+directory+'\\'+SS_msg+'.png')
        except:
            pass

def word_cov(data):
        try:
            dirr= "temp"
            doc = docx.Document()
            doc.add_heading('Evidence', 0)
            dummy=len(mylist)
            for i in range(dummy):
                doc.add_heading(mylist[i], 5)
                doc.add_picture('D:\End to End\Evidence\\'+dirr+'\\'+mylist[i]+'.png', width=Inches(6.5), height=Inches(3.5))
            doc.save('D:\End to End\Evidence\\'+data+'.docx')
            parent1 = "D:/End to End/Evidence/"
            path2 = os.path.join(parent1, dirr)
            shutil.rmtree(path2)
            mylist.clear()
        except:
            pass

def DB_Update(value,data):
     conn=db.connect(host='localhost',user='root',password='Muruga@2000')
     cursor=conn.cursor()
     cursor.execute("UPDATE pydatabase.py_automate SET Result='"+value+"' WHERE Scenario='"+data+"' and Run_Indicator='Y';")  
     conn.commit()
     cursor.close()
     conn.close()

#Y=input_data(Db_Connection())

"""
for i in range (len(Y)):
    print(Y['NAME'][i])
    print(Y['PHONE'][i])
    print(Y['ADDRESS'][i])

"""
















"""
i=0
df=pd.read_excel('dataframe.xlsx',dtype=str)
#print(df['Name'][0])

def dictget():
    Name=[]
    Phone=[]
    Address=[]
    for i in range(df.shape[0]):
        Name.append(df['Name'][i])
        Phone.append(df['Phone'][i])
        Address.append(df['Address'][i])

    mydict={"NAME":Name,"PHONE":Phone,"ADDRESS":Address}

    return mydict
#print(Name)
#print(Phone)
#print(Address)

#print(mydict)

#print(mydict["NAME"][1])

#print(mydict['NAME'][0])

A=dictget()
#print(A["NAME"][0])
#B=len(A["NAME"])
for i in range (len(A["NAME"])):
    print(A["NAME"][i])
    print(A["PHONE"][i])
    print(A["ADDRESS"][i])

#print()

"""